from PyPDF2 import PdfReader
from docx import Document
import pytesseract
from PIL import Image
import os
import base64
import unicodedata
import cv2


def extract_pdf(pdf_path):
    """Extract form fields from a PDF file."""
    reader = PdfReader(pdf_path)
    fields = reader.get_fields()
    return {k: v.get("/V", None) for k, v in fields.items()}


def parse_docx(doc_path):
    """Extract personal information from a formatted Word document."""
    doc = Document(doc_path)

    # Extract personal details from first table
    table = doc.tables[1]
    extract = lambda table, row, col: table.rows[row].cells[col].text.strip()

    data = {
        "last_name": extract(table, 0, 2),
        "first_middle_names": extract(table, 1, 2),
        "address": extract(table, 2, 2),
        "country_of_domicile": extract(table, 3, 2),
        "date_of_birth": extract(table, 4, 2),
        "nationality": extract(table, 5, 2),
        "id_passport_number": extract(table, 6, 2),
        "id_type": extract(table, 7, 2),
        "id_issue_date": extract(table, 8, 2),
        "id_expiry_date": extract(table, 9, 2),
    }

    # Determine gender from checkbox
    gender_cell = extract(table, 10, 2)
    if "☒ Male" in gender_cell:
        data["gender"] = "male"
    elif "☒ Female" in gender_cell:
        data["gender"] = "female"
    else:
        data["gender"] = "unknown"

    # Extract contact information from second table
    table = doc.tables[3]
    data.update(
        {
            "telephone": "".join(extract(table, 0, 2).split()[1:]),
            "email": extract(table, 2, 2).split()[1],
        }
    )

    return data


def normalize_text(input_text):
    """Normalize Unicode text by removing diacritics and special characters."""
    nfkd_form = unicodedata.normalize("NFD", input_text)
    nfkd_form = nfkd_form.replace("ł", "l").replace("Ł", "L")
    return "".join([char for char in nfkd_form if unicodedata.category(char) != "Mn"])


def preprocess_image(image_path):
    """Enhance image for better OCR performance."""
    img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
    img = cv2.resize(img, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
    img = cv2.adaptiveThreshold(
        img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, blockSize=9, C=15
    )
    return img


def extract_text(image_path):
    """Extract text from an image using OCR."""
    img = preprocess_image(image_path)
    text = pytesseract.image_to_string(img, config="--psm 11")
    return text


def save_passport_image(json_data, output_dir, filename="passport.png"):
    """Save base64-encoded passport image to file."""
    if "passport" not in json_data:
        raise ValueError("Field 'passport' not found in the provided JSON data.")

    os.makedirs(output_dir, exist_ok=True)
    image_data = base64.b64decode(json_data["passport"])
    filepath = os.path.join(output_dir, filename)

    with open(filepath, "wb") as f:
        f.write(image_data)

    return filepath


def save_docx_file(json_data, output_dir, filename="profile.docx"):
    """Save base64-encoded Word document to file."""
    if "profile" not in json_data:
        raise ValueError("Field 'profile' not found in the provided JSON data.")

    os.makedirs(output_dir, exist_ok=True)
    docx_data = base64.b64decode(json_data["profile"])
    filepath = os.path.join(output_dir, filename)

    with open(filepath, "wb") as f:
        f.write(docx_data)

    return filepath


def save_pdf_file(json_data, output_dir, filename="account.pdf"):
    """Save base64-encoded PDF to file."""
    if "account" not in json_data:
        raise ValueError("Field 'account' not found in the provided JSON data.")

    os.makedirs(output_dir, exist_ok=True)
    pdf_data = base64.b64decode(json_data["account"])
    filepath = os.path.join(output_dir, filename)

    with open(filepath, "wb") as f:
        f.write(pdf_data)

    return filepath


def save_description_txt(json_data, output_dir, filename="description.txt"):
    """Save base64-encoded text description to file."""
    if "description" not in json_data:
        raise ValueError("Field 'description' not found in the provided JSON data.")

    os.makedirs(output_dir, exist_ok=True)
    txt_data = base64.b64decode(json_data["description"])
    filepath = os.path.join(output_dir, filename)

    with open(filepath, "wb") as f:
        f.write(txt_data)

    return filepath
